"""nodes.document_parsers
Document parsing utilities for extracting text from PDF and DOCX files.

This module provides parsers for different document formats that can be used
for style extraction. Each parser extracts text content and splits it into
sentences for further processing.
"""

import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
import re

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import nltk
    from nltk.tokenize import sent_tokenize
except ImportError:
    nltk = None
    sent_tokenize = None

logger = logging.getLogger(__name__)


def _ensure_nltk_data():
    """Ensure NLTK punkt tokenizer is available."""
    if nltk is None:
        raise ImportError("NLTK is required for sentence tokenization. Install with: pip install nltk")
    
    # Try to find the tokenizer data, handling both old and new NLTK versions
    try:
        # Try the newer punkt_tab first (NLTK 3.8+)
        nltk.data.find('tokenizers/punkt_tab')
    except LookupError:
        try:
            # Try the older punkt format
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            # Download both to ensure compatibility
            logger.info("Downloading NLTK punkt tokenizer...")
            try:
                nltk.download('punkt_tab', quiet=True)
            except Exception:
                # Fall back to older punkt if punkt_tab fails
                nltk.download('punkt', quiet=True)


def _basic_sentence_split(text: str) -> List[str]:
    """Basic sentence splitting fallback when NLTK is not available."""
    # Simple regex-based sentence splitting
    sentences = re.split(r'[.!?]+\s+', text)
    # Remove any trailing punctuation from the last sentence
    if sentences and sentences[-1]:
        sentences[-1] = re.sub(r'[.!?]+$', '', sentences[-1])
    return [s.strip() for s in sentences if s.strip()]


def split_into_sentences(text: str, language: str = "english") -> List[str]:
    """Split text into sentences using NLTK or basic regex fallback.
    
    Args:
        text: The input text to split
        language: Language for NLTK tokenizer (default: english)
    
    Returns:
        List of sentences
    """
    if not text.strip():
        return []
    
    # Try NLTK first
    if sent_tokenize is not None:
        try:
            _ensure_nltk_data()
            sentences = sent_tokenize(text, language=language)
            return [s.strip() for s in sentences if s.strip()]
        except Exception as e:
            logger.warning(f"NLTK sentence tokenization failed: {e}. Using basic fallback.")
    
    # Fallback to basic regex splitting
    return _basic_sentence_split(text)


def parse_pdf(file_path: str, language: str = "english") -> List[str]:
    """Parse PDF file and extract sentences.
    
    Args:
        file_path: Path to the PDF file
        language: Language for sentence tokenization
    
    Returns:
        List of sentences extracted from the PDF
    
    Raises:
        ImportError: If PyPDF2 is not installed
        FileNotFoundError: If the PDF file doesn't exist
        ValueError: If the PDF cannot be parsed
    """
    if pypdf is None:
        raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")
    
    pdf_path = Path(file_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    logger.info(f"Parsing PDF file: {file_path}")
    
    text_content = ""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
    
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file {file_path}: {e}")
    
    if not text_content.strip():
        raise ValueError(f"No text content found in PDF file: {file_path}")
    
    # Clean up the text
    text_content = re.sub(r'\s+', ' ', text_content)  # Normalize whitespace
    text_content = text_content.strip()
    
    # Split into sentences
    sentences = split_into_sentences(text_content, language)
    
    logger.info(f"Extracted {len(sentences)} sentences from PDF")
    return sentences


def parse_docx(file_path: str, language: str = "english") -> List[str]:
    """Parse DOCX file and extract sentences.
    
    Args:
        file_path: Path to the DOCX file
        language: Language for sentence tokenization
    
    Returns:
        List of sentences extracted from the DOCX
    
    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If the DOCX file doesn't exist
        ValueError: If the DOCX cannot be parsed
    """
    if Document is None:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    docx_path = Path(file_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    logger.info(f"Parsing DOCX file: {file_path}")
    
    text_content = ""
    try:
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " "
            text_content += "\n"
    
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file {file_path}: {e}")
    
    if not text_content.strip():
        raise ValueError(f"No text content found in DOCX file: {file_path}")
    
    # Clean up the text
    text_content = re.sub(r'\s+', ' ', text_content)  # Normalize whitespace
    text_content = text_content.strip()
    
    # Split into sentences
    sentences = split_into_sentences(text_content, language)
    
    logger.info(f"Extracted {len(sentences)} sentences from DOCX")
    return sentences


def parse_doc(file_path: str, language: str = "english") -> List[str]:
    """Parse DOC file and extract sentences.
    
    Note: This is a basic implementation that attempts to parse DOC files
    using python-docx. For better DOC support, consider using additional
    libraries like python-docx2txt or msoffcrypto-tool.
    
    Args:
        file_path: Path to the DOC file
        language: Language for sentence tokenization
    
    Returns:
        List of sentences extracted from the DOC
    
    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If the DOC file doesn't exist
        ValueError: If the DOC cannot be parsed
    """
    # Try to parse DOC file with python-docx (limited support)
    try:
        return parse_docx(file_path, language)
    except Exception as e:
        raise ValueError(f"Failed to parse DOC file {file_path}. DOC files have limited support. "
                        f"Consider converting to DOCX format. Error: {e}")


def parse_document(file_path: str, file_type: str, language: str = "english") -> List[str]:
    """Parse a document file based on its type.
    
    Args:
        file_path: Path to the document file
        file_type: Type of document ('pdf', 'docx', 'doc')
        language: Language for sentence tokenization
    
    Returns:
        List of sentences extracted from the document
    
    Raises:
        ValueError: If file_type is not supported
        Various exceptions from specific parsers
    """
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return parse_pdf(file_path, language)
    elif file_type == 'docx':
        return parse_docx(file_path, language)
    elif file_type == 'doc':
        return parse_doc(file_path, language)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: pdf, docx, doc")


def create_document_entries(sentences: List[str], source_language: str) -> List[Dict[str, Any]]:
    """Create TMX-like entries from document sentences for style guide generation.
    
    Args:
        sentences: List of sentences from the document
        source_language: Language of the document
    
    Returns:
        List of entry dictionaries compatible with TMX processing
    """
    entries = []
    
    for i, sentence in enumerate(sentences):
        if not sentence.strip():
            continue
            
        # Create a TMX-like entry structure
        entry = {
            "source": sentence,
            "target": sentence,  # For style extraction, we use the same text
            "source_lang": source_language,
            "target_lang": source_language,
            "usage_count": 1,  # All sentences have equal weight initially
            "document_sentence": True,  # Flag to indicate this is from document parsing
            "sentence_index": i
        }
        entries.append(entry)
    
    logger.info(f"Created {len(entries)} document entries from {len(sentences)} sentences")
    return entries